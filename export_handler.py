import csv
import json
import os
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment
from copy import copy


def export_rfp(filepath, questions, rfp_id):
    ext = filepath.rsplit(".", 1)[-1].lower()
    os.makedirs("exports", exist_ok=True)

    if ext == "csv":
        return _export_csv(filepath, questions, rfp_id)
    elif ext in ("xlsx", "xls"):
        return _export_xlsx(filepath, questions, rfp_id)
    elif ext == "docx":
        return _export_docx(filepath, questions, rfp_id)
    raise ValueError(f"Unsupported file type: {ext}")


def _build_lookup(questions):
    """Map first 120 chars of question text (lowercased) to question dict."""
    lookup = {}
    for q in questions:
        if q.get("status") == "answered" and q.get("answer"):
            key = str(q["question_text"])[:120].lower().strip()
            lookup[key] = q
    return lookup


def _export_xlsx(filepath, questions, rfp_id):
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active

    # Locate header row (first row with content)
    header_row_idx = None
    headers = []
    for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
        if any(c for c in row if c is not None):
            header_row_idx = i
            headers = [str(c).strip().lower() if c is not None else "" for c in row]
            break

    if header_row_idx is None:
        raise ValueError("Could not find header row")

    # Detect relevant column indices (1-based)
    req_col = resp_col = comment_col = None
    for j, h in enumerate(headers):
        if any(k in h for k in ("requirement", "criteria", "question", "description")):
            if req_col is None:
                req_col = j + 1
        if any(k in h for k in ("vendor response", "response required", "vendor resp")):
            resp_col = j + 1
        if any(k in h for k in ("comment", "assumption", "notes", "vendor comment")):
            comment_col = j + 1

    # Fallback column positions
    max_col = ws.max_column
    if resp_col is None:
        resp_col = max_col - 1 if max_col > 2 else max_col
    if comment_col is None:
        comment_col = max_col

    lookup = _build_lookup(questions)

    green_fill = PatternFill(start_color="00BF6F", end_color="00BF6F", fill_type="solid")
    amber_fill = PatternFill(start_color="F5A623", end_color="F5A623", fill_type="solid")

    filled = 0
    for row_num in range(header_row_idx + 1, ws.max_row + 1):
        if req_col is None:
            break
        cell_val = ws.cell(row=row_num, column=req_col).value
        if not cell_val:
            continue
        key = str(cell_val)[:120].lower().strip()
        if key not in lookup:
            continue

        q = lookup[key]
        rc = q.get("response_code", "")
        answer = q.get("answer", "")
        # Strip internal review notes from export
        answer = answer.split("\n\n⚠")[0].strip()

        resp_cell = ws.cell(row=row_num, column=resp_col)
        resp_cell.value = rc
        resp_cell.font = Font(bold=True)
        resp_cell.alignment = Alignment(horizontal="center")
        if rc == "F":
            resp_cell.fill = green_fill
        elif rc in ("P", "C"):
            resp_cell.fill = amber_fill

        comment_cell = ws.cell(row=row_num, column=comment_col)
        comment_cell.value = answer
        comment_cell.alignment = Alignment(wrap_text=True, vertical="top")
        filled += 1

    export_path = os.path.join("exports", f"rfp_{rfp_id}_naughtrfp_export.xlsx")
    wb.save(export_path)
    return export_path


def _export_csv(filepath, questions, rfp_id):
    with open(filepath, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        fieldnames = list(reader.fieldnames or [])

    if "Vendor Response" not in fieldnames:
        fieldnames += ["Vendor Response", "Vendor Comments"]

    lookup = _build_lookup(questions)

    for row in rows:
        for val in row.values():
            key = str(val)[:120].lower().strip() if val else ""
            if key in lookup:
                q = lookup[key]
                answer = q.get("answer", "")
                answer = answer.split("\n\n⚠")[0].strip()
                row["Vendor Response"] = q.get("response_code", "")
                row["Vendor Comments"] = answer
                break

    export_path = os.path.join("exports", f"rfp_{rfp_id}_naughtrfp_export.csv")
    with open(export_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)

    return export_path


def _export_docx(filepath, questions, rfp_id):
    """
    Export RFP answers back into a Word (.docx) document.

    The exported document contains a table with columns mirroring the original
    structure plus Vendor Response and Vendor Comments columns.  Response codes
    are colour-coded using Word cell shading:
        F  → green  (#00BF6F)
        P  → amber  (#F5A623)
        C  → blue   (#4A90D9)
        NE → red    (#D0021B)
        N  → red    (#D0021B)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_shading(cell, hex_color: str):
        """Apply solid background shading to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        tcPr.append(shd)

    _CODE_COLORS = {
        "F":  "00BF6F",  # green
        "P":  "F5A623",  # amber
        "C":  "4A90D9",  # blue
        "NE": "D0021B",  # red
        "N":  "D0021B",  # red
    }

    lookup = _build_lookup(questions)

    # ── Try to preserve original table structure ──────────────────────────────
    # Open the source document and collect table data; we'll rebuild the table
    # in the export document with answer columns appended.
    src_doc = Document(filepath)

    export_doc = Document()
    export_doc.add_heading("NaughtRFP — Vendor Response Export", level=1)

    if src_doc.tables:
        for src_table in src_doc.tables:
            if not src_table.rows:
                continue

            src_headers = [cell.text.strip() for cell in src_table.rows[0].cells]

            # Build export column list: original columns + response/comments if absent
            export_headers = list(src_headers)
            resp_col_name = "Vendor Response"
            comments_col_name = "Vendor Comments"
            if resp_col_name not in export_headers:
                export_headers.append(resp_col_name)
            if comments_col_name not in export_headers:
                export_headers.append(comments_col_name)

            resp_idx     = export_headers.index(resp_col_name)
            comments_idx = export_headers.index(comments_col_name)
            n_cols       = len(export_headers)

            # Create table in export doc
            tbl = export_doc.add_table(rows=1, cols=n_cols)
            tbl.style = "Table Grid"

            # Header row
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(export_headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs
                if run:
                    run[0].bold = True

            # Data rows
            for src_row in src_table.rows[1:]:
                src_cells = [cell.text.strip() for cell in src_row.cells]
                if not any(src_cells):
                    continue

                # Try to match the row to an answered question
                q_match = None
                for cell_text in src_cells:
                    key = cell_text[:120].lower().strip()
                    if key in lookup:
                        q_match = lookup[key]
                        break

                new_row = tbl.add_row().cells

                # Copy original columns
                for i, h in enumerate(export_headers):
                    if i < len(src_headers) and i < len(src_cells):
                        new_row[i].text = src_cells[i]

                # Fill response + comments
                if q_match:
                    rc     = q_match.get("response_code", "")
                    answer = q_match.get("answer", "")
                    answer = answer.split("\n\n⚠")[0].strip()

                    new_row[resp_idx].text     = rc
                    new_row[comments_idx].text = answer

                    # Bold + centre the response code cell
                    para = new_row[resp_idx].paragraphs[0]
                    para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER = 1
                    if para.runs:
                        para.runs[0].bold = True

                    # Colour-code the response code cell
                    hex_color = _CODE_COLORS.get(rc.upper(), "")
                    if hex_color:
                        _set_cell_shading(new_row[resp_idx], hex_color)

            export_doc.add_paragraph()  # spacing between tables

    else:
        # No tables in source: create a simple two-column answer table
        answered = [q for q in questions if q.get("status") == "answered" and q.get("answer")]
        if answered:
            tbl = export_doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            for i, h in enumerate(["Requirement", "Vendor Response", "Vendor Comments"]):
                hdr[i].text = h
                if hdr[i].paragraphs[0].runs:
                    hdr[i].paragraphs[0].runs[0].bold = True

            for q in answered:
                rc     = q.get("response_code", "")
                answer = q.get("answer", "")
                answer = answer.split("\n\n⚠")[0].strip()

                row_cells = tbl.add_row().cells
                row_cells[0].text = q.get("question_text", "")
                row_cells[1].text = rc
                row_cells[2].text = answer

                para = row_cells[1].paragraphs[0]
                para.alignment = 1
                if para.runs:
                    para.runs[0].bold = True
                hex_color = _CODE_COLORS.get(rc.upper(), "")
                if hex_color:
                    _set_cell_shading(row_cells[1], hex_color)

    export_path = os.path.join("exports", f"rfp_{rfp_id}_naughtrfp_export.docx")
    export_doc.save(export_path)
    return export_path
