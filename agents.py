import anthropic
import csv
import httpx
import json
import os
import re
import time
import openpyxl
from concurrent.futures import ThreadPoolExecutor, as_completed


_LITELLM_BASE_URL: str | None = None
_WEB_SEARCH_ENABLED: bool = True


def _parse_docx_rows(filepath: str) -> list[dict]:
    """
    Parse a .docx file into a list of row dicts compatible with the Parser Agent pipeline.

    Strategy (in priority order):
    1. Tables: iterate every table, treating the first row as headers if it looks like one,
       otherwise using positional column names (col_0, col_1, …).  Each subsequent row
       becomes a dict keyed by those headers.
    2. Smart paragraph extraction (fallback): if no table rows are found, apply heuristics
       to identify likely requirement paragraphs using paragraph style names and content
       patterns.  Headings become category markers; numbered/list items and long Normal
       paragraphs become requirement rows.
    3. Last-resort full-text extraction: if still fewer than 3 rows, extract ALL text and
       return any chunk longer than 50 chars — better than returning nothing.

    Returns structured dicts: {"category": str, "requirement": str, "seq": int} when
    headings are detected, or {"requirement": str, "seq": int} for unheaded content,
    or {"text": str} for table-sourced rows (preserving backward compat).

    The downstream header-detection logic in _parse_rfp is unchanged — the Parser Agent's
    Claude call still determines which column holds the requirement text.
    """
    try:
        from docx import Document
    except ImportError:
        return []

    # Numbered-item pattern: "1.", "1)", "a.", "a)", "(1)", "(a)" at start of text
    _NUMBERED_RE = re.compile(r'^(\(?\d+\)?\.?\)?\s|\(?[a-zA-Z]\)?\.?\)?\s)', re.UNICODE)

    rows: list[dict] = []
    try:
        doc = Document(filepath)

        # ── Strategy 1: Tables ─────────────────────────────────────────────────
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(cells)

            if not table_rows:
                continue

            # Detect whether first row is a header (heuristic: cells are short, no digits-only)
            first = table_rows[0]
            looks_like_header = all(
                len(c) < 80 and not c.isdigit() for c in first if c
            )

            if looks_like_header and len(table_rows) > 1:
                headers = [c if c else f"col_{i}" for i, c in enumerate(first)]
                data_rows = table_rows[1:]
            else:
                headers = [f"col_{i}" for i in range(len(first))]
                data_rows = table_rows

            for cells in data_rows:
                if not any(c for c in cells):
                    continue
                row_dict = {
                    headers[i]: cells[i] if i < len(cells) else ""
                    for i in range(len(headers))
                }
                rows.append(row_dict)

        # ── Strategy 2: Smart paragraph extraction ────────────────────────────
        if not rows:
            current_category = "General"
            seq = 0
            has_headings = False

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"

                # Heading styles → category markers
                if style_name.startswith("Heading"):
                    current_category = text
                    has_headings = True
                    continue

                # List styles → strong signal for requirement items
                is_list_style = any(s in style_name for s in ("List", "Bullet", "Number"))

                # Numbered/lettered item pattern
                is_numbered = bool(_NUMBERED_RE.match(text))

                # Length filter for Normal text
                is_long_enough = len(text) >= 40

                # Skip likely titles, footers, single words
                is_noise = (
                    len(text) < 15 or
                    text.isupper() and len(text) < 40 or   # ALL CAPS short labels
                    "\n" not in text and len(text.split()) <= 2  # single/double word
                )

                if is_noise and not is_list_style and not is_numbered:
                    continue

                if is_list_style or is_numbered or is_long_enough:
                    seq += 1
                    if has_headings:
                        rows.append({
                            "category": current_category,
                            "requirement": text,
                            "seq": seq,
                        })
                    else:
                        rows.append({
                            "requirement": text,
                            "seq": seq,
                        })

        # ── Strategy 3: Last-resort full-text extraction ──────────────────────
        if len(rows) < 3:
            all_text_parts = []

            # All paragraphs
            for para in doc.paragraphs:
                t = para.text.strip()
                if t and len(t) > 50:
                    all_text_parts.append(t)

            # All table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and len(t) > 50:
                            all_text_parts.append(t)

            # Deduplicate while preserving order
            seen_texts: set[str] = set()
            for chunk in all_text_parts:
                if chunk not in seen_texts:
                    seen_texts.add(chunk)
                    rows.append({"text": chunk})

    except Exception:
        pass

    return rows

# In-process page cache: url → (fetched_at_epoch, content)
_PAGE_CACHE: dict[str, tuple[float, str]] = {}
_PAGE_CACHE_TTL: int = 3600

# Model assignments — change here to update all agents at once
# Sonnet: heavy reasoning, tool use, multi-step decisions
# Haiku:  fast extraction, structured output, pattern matching
_MODEL      = "claude-sonnet-4-6"   # Analysis, Answer, Demo Prep, AI KB Search
_MODEL_FAST = "claude-haiku-4-5"    # Customer, Parser, Web Summarizer
_ANSWER_WORKERS = 6


def set_web_search_enabled(enabled: bool) -> None:
    global _WEB_SEARCH_ENABLED
    _WEB_SEARCH_ENABLED = enabled


def set_litellm_base_url(url: str | None) -> None:
    global _LITELLM_BASE_URL
    _LITELLM_BASE_URL = url or None


def _make_client(api_key: str) -> anthropic.Anthropic:
    """Create Anthropic client routing through LiteLLM proxy with SSL bypass for corporate environments."""
    import warnings
    warnings.filterwarnings("ignore", message=".*verify=False.*")
    warnings.filterwarnings("ignore", message=".*Unverified HTTPS.*")
    kwargs = dict(
        api_key=api_key,
        http_client=httpx.Client(verify=False, timeout=120.0),
    )
    if _LITELLM_BASE_URL:
        kwargs["base_url"] = _LITELLM_BASE_URL
    return anthropic.Anthropic(**kwargs)


OKTA_KNOWLEDGE = """
Okta is the world's leading Identity platform. Key products for IGA/IAM RFPs:

OKTA IDENTITY GOVERNANCE (OIG)
- Access certifications / access reviews (automated, AI-assisted)
- Entitlement management and access requests
- Separation of Duties (SoD) policy detection and enforcement
- Identity lifecycle management (Joiner-Mover-Leaver automation)
- Risk-based access decisions
- Comprehensive audit trails and compliance reporting
- AI-powered access recommendations and anomaly detection
- Role management and RBAC/PBAC support

OKTA LIFECYCLE MANAGEMENT (LCM)
- HR-driven Joiner-Mover-Leaver automation
- Automated provisioning and deprovisioning to 200+ apps
- Universal Directory as the master identity store
- Custom attribute mapping and transformation
- Multi-source identity aggregation
- Non-employee / contractor / third-party identity management

OKTA WORKFLOWS
- No-code/low-code automation for complex identity processes
- Custom provisioning logic with 30+ connectors
- Integration with ITSM, HR, and custom systems
- Event-driven triggers for identity events

OKTA ACCESS MANAGEMENT
- SSO via SAML 2.0, OIDC, WS-Fed
- Adaptive MFA with 20+ factor types
- Passwordless authentication (FIDO2/WebAuthn)
- Risk-based and step-up authentication
- API Access Management (OAuth 2.0)

OKTA PRIVILEGED ACCESS (PAM)
- Privileged account vaulting
- Just-in-time access provisioning
- Session recording and monitoring
- Infrastructure access management

PLATFORM & COMPLIANCE
- Cloud-native SaaS, 99.99% uptime SLA
- Data residency: US, EU, Canada (PIPEDA compliant)
- FedRAMP Authorized (High) for US Federal
- SOC 2 Type II, ISO 27001, ISO 27018 certified
- GDPR and CCPA compliant
- Okta Integration Network: 7,000+ pre-built integrations

AI CAPABILITIES (Okta AI)
- AI-generated entitlement descriptions
- Intelligent access review recommendations
- Anomaly detection and risk scoring
- AI agent identity support (NHI management)
- Predictive access governance

INTEGRATIONS
- HR: Workday, BambooHR, SAP SuccessFactors, ADP, UKG
- ITSM: ServiceNow, Jira Service Management
- Directories: Active Directory, LDAP, Azure AD
- Cloud: AWS, Azure, GCP
- SCIM 2.0 provisioning standard
- REST API and webhooks for custom integrations
"""

SEARCH_WEB_TOOL = {
    "name": "search_web",
    "description": (
        "Search the web for current Okta product information, compliance details, "
        "certification status, uptime data, and technical documentation. "
        "Always search okta.com and trust.okta.com sources. "
        "Use when the knowledge base doesn't have enough context or you need up-to-date facts."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "query": {
                "type": "string",
                "description": "Specific search query, e.g. 'Okta FedRAMP High authorization' or 'Okta 99.99 uptime SLA trust portal'"
            }
        },
        "required": ["query"]
    }
}

SEARCH_KB_TOOL = {
    "name": "search_knowledge_base",
    "description": "Search the knowledge base for relevant past RFP answers. Call this before generating an answer.",
    "input_schema": {
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "Search query based on the requirement topic"}
        },
        "required": ["query"]
    }
}

FLAG_REVIEW_TOOL = {
    "name": "flag_for_review",
    "description": "Flag this requirement for human review. Use when confidence is below 60%, or when the requirement involves specific pricing, SLA commitments, legal obligations, or highly custom scenarios.",
    "input_schema": {
        "type": "object",
        "properties": {
            "reason": {"type": "string", "description": "Why this requires human review"}
        },
        "required": ["reason"]
    }
}


_OKTA_PAGES = {
    "compliance":     "https://trust.okta.com/compliance",
    "security":       "https://trust.okta.com/security",
    "trust":          "https://trust.okta.com",
    "iga":            "https://www.okta.com/products/identity-governance/",
    "lifecycle":      "https://www.okta.com/products/lifecycle-management/",
    "mfa":            "https://www.okta.com/products/adaptive-multi-factor-authentication/",
    "sso":            "https://www.okta.com/products/single-sign-on/",
    "workflows":      "https://www.okta.com/products/okta-workflows/",
    "pam":            "https://www.okta.com/products/privileged-access/",
    "integrations":   "https://www.okta.com/integrations/",
    "ai":             "https://www.okta.com/products/okta-ai/",
    "certifications": "https://trust.okta.com/compliance",
}


def _fetch_page(url: str) -> str:
    """Fetch a URL, strip HTML, cache result for _PAGE_CACHE_TTL seconds."""
    import warnings
    warnings.filterwarnings("ignore")

    cached = _PAGE_CACHE.get(url)
    if cached and (time.time() - cached[0]) < _PAGE_CACHE_TTL:
        return cached[1]

    try:
        r = httpx.get(url, verify=False, timeout=10, follow_redirects=True,
                      headers={"User-Agent": "Mozilla/5.0 NaughtRFP/1.0"})
        if r.status_code != 200:
            return ""
        text = re.sub(r"<script[^>]*>.*?</script>", "", r.text, flags=re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>",  "", text,   flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()[:4000]
    except Exception:
        text = ""

    _PAGE_CACHE[url] = (time.time(), text)
    return text


def _ddg_search(query: str) -> str:
    """DuckDuckGo Instant Answer API — no key required."""
    import urllib.parse, warnings
    warnings.filterwarnings("ignore")
    try:
        q = urllib.parse.quote_plus(f"site:okta.com OR site:trust.okta.com {query}")
        r = httpx.get(
            f"https://api.duckduckgo.com/?q={q}&format=json&no_html=1&skip_disambig=1",
            verify=False, timeout=8, follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 NaughtRFP/1.0"},
        )
        if r.status_code == 200:
            d = r.json()
            parts = []
            if d.get("AbstractText"):
                parts.append(d["AbstractText"])
            for topic in d.get("RelatedTopics", [])[:4]:
                if isinstance(topic, dict) and topic.get("Text"):
                    parts.append(topic["Text"])
            return "\n".join(parts) if parts else ""
    except Exception:
        pass
    return ""


def _do_web_search(query: str, api_key: str, base_url: str | None) -> str:
    """
    Multi-source web search:
    1. DuckDuckGo Instant Answer (free, no key)
    2. Targeted Okta page fetch based on query keywords
    Then summarise with Claude.
    """
    # Collect raw content
    raw_parts = []

    ddg = _ddg_search(query)
    if ddg:
        raw_parts.append(f"[DuckDuckGo]\n{ddg}")

    # Pick most relevant Okta pages by keyword match
    ql = query.lower()
    for kw, url in _OKTA_PAGES.items():
        if kw in ql or any(w in ql for w in kw.split()):
            page = _fetch_page(url)
            if page:
                raw_parts.append(f"[{url}]\n{page[:1500]}")
            break

    if not raw_parts:
        # Fallback: fetch trust portal
        page = _fetch_page("https://trust.okta.com")
        if page:
            raw_parts.append(f"[trust.okta.com]\n{page[:2000]}")

    if not raw_parts:
        return "[Web search returned no results]"

    combined = "\n\n".join(raw_parts)[:5000]

    # Summarise with Claude
    try:
        client = _make_client(api_key)
        resp = client.messages.create(
            model=_MODEL_FAST,
            max_tokens=400,
            system="You extract key facts from web content to help answer enterprise RFP questions about Okta.",
            messages=[{
                "role": "user",
                "content": (
                    f"Query: {query}\n\n"
                    f"Web content:\n{combined}\n\n"
                    "Extract the 3-5 most relevant facts for answering an enterprise RFP. "
                    "Be specific — include numbers, certification names, SLA percentages. "
                    "Format as bullet points."
                )
            }],
        )
        return resp.content[0].text
    except Exception as e:
        return combined[:800]


def _record(db, rfp_id, agent, resp):
    """Record token usage from an Anthropic response object."""
    try:
        u = resp.usage
        db.record_tokens(rfp_id, agent, resp.model,
                         getattr(u, "input_tokens", 0),
                         getattr(u, "output_tokens", 0))
    except Exception:
        pass


def _extract_json(text, kind="object"):
    """Extract first JSON object or array from text, handling markdown code fences."""
    if not text:
        return None

    # Strip markdown code fences first
    cleaned = re.sub(r'```(?:json)?\s*', '', text)
    cleaned = re.sub(r'```', '', cleaned).strip()

    # Try parsing the whole thing first
    try:
        return json.loads(cleaned)
    except Exception:
        pass

    # Fall back to regex extraction
    pattern = r'\{.*\}' if kind == "object" else r'\[.*\]'
    match = re.search(pattern, cleaned, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass

    # Last resort: try original text with regex
    match = re.search(pattern, text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass

    return None


class AgentPipeline:
    def __init__(self, api_key, db, event_queue):
        self.api_key = api_key
        self.db = db
        self.q = event_queue
        self.client = _make_client(api_key)

    def emit(self, event_type, agent=None, message="", data=None):
        self.q.put({
            "type": event_type,
            "agent": agent,
            "message": message,
            "timestamp": time.time(),
            **({"data": data} if data else {}),
        })

    def process_rfp(self, rfp_id, filepath):
        self.db.update_rfp(rfp_id, status="processing")

        try:
            # ── 0. Customer Detection (if not already done on upload) ─────────
            rfp = self.db.get_rfp(rfp_id)
            if not rfp.get("customer_info"):
                self.emit("agent_start", "Customer Agent", "Identifying who issued this RFP...")
                try:
                    ci = detect_customer(self.api_key, filepath,
                                        rfp.get("name", "rfp"))
                    self.db.update_rfp(rfp_id, customer_info=json.dumps(ci))
                    self.emit("agent_complete", "Customer Agent",
                              f"Identified: {ci.get('customer_name', 'Unknown')} — {ci.get('project_name', '')}")
                except Exception as e:
                    self.emit("agent_progress", "Customer Agent", f"Detection skipped: {e}")

            # ── 1. Parser Agent ───────────────────────────────────────────────
            self.emit("agent_start", "Parser Agent", "Reading RFP and detecting column structure...")
            questions = self._parse_rfp(rfp_id, filepath)
            if not questions:
                self.db.update_rfp(rfp_id, status="error")
                fname = os.path.basename(filepath)
                ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                if ext == "docx":
                    self.emit("error", message=(
                        f"No requirements could be extracted from {fname}. "
                        "For Word documents, NaughtRFP works best with documents that contain "
                        "tables or numbered requirement lists. "
                        "Try converting to CSV or XLSX for best results."
                    ))
                else:
                    self.emit("error", message=f"No requirements found in {fname}.")
                return
            self.emit("agent_complete", "Parser Agent",
                      f"Extracted {len(questions)} requirements across {len(set(q['category'] for q in questions))} categories",
                      {"count": len(questions)})

            # ── 2. Analysis Agent ─────────────────────────────────────────────
            self.emit("agent_start", "Analysis Agent",
                      f"Mapping {len(questions)} requirements to Okta products and risk levels...")
            questions = self._analyze_questions(rfp_id, questions)
            self.emit("agent_complete", "Analysis Agent",
                      "Requirements categorized and mapped to Okta product areas")

            # ── 3. Research Agent: bulk KB pre-fetch ──────────────────────────
            self.emit("agent_start", "Research Agent",
                      f"Pre-fetching KB context for all {len(questions)} requirements...")
            kb_cache: dict[int, list] = {}
            for q in questions:
                hits = self.db.search_knowledge_base(q["question_text"][:150], limit=3)
                if hits:
                    kb_cache[q["id"]] = hits
            kb_hits = sum(1 for v in kb_cache.values() if v)
            self.emit("agent_complete", "Research Agent",
                      f"KB pre-fetch complete — {kb_hits}/{len(questions)} questions have matching context")

            # ── 4. Answer Agent (parallel, {_ANSWER_WORKERS} workers) ─────────
            self.emit("agent_start", "Answer Agent",
                      f"Processing {len(questions)} requirements ({_ANSWER_WORKERS} parallel workers)...")

            answered = flagged = 0
            fit_scores = []
            risk_scores = []
            completed = [0]

            def _process_one(args):
                idx, q = args
                pre_ctx = kb_cache.get(q["id"], [])
                return idx, q, self._research_and_answer(rfp_id, q, pre_context=pre_ctx)

            with ThreadPoolExecutor(max_workers=_ANSWER_WORKERS) as pool:
                futures = {pool.submit(_process_one, (i, q)): i
                           for i, q in enumerate(questions)}
                for future in as_completed(futures):
                    try:
                        idx, q, result = future.result()
                    except Exception as e:
                        self.emit("agent_progress", "Answer Agent", f"  ✗ Error: {e}")
                        flagged += 1
                        continue

                    completed[0] += 1
                    pct = completed[0]
                    total = len(questions)

                    if result.get("flagged"):
                        flagged += 1
                        self.emit("agent_progress", "Answer Agent",
                                  f"  [{pct}/{total}] ⚑ {result.get('review_reason','')[:70]}",
                                  {"current": pct, "total": total})
                    else:
                        answered += 1
                        fit_scores.append(result.get("fit_score", 3))
                        risk_scores.append(result.get("risk_score", 3))
                        self.emit("agent_progress", "Answer Agent",
                                  f"  [{pct}/{total}] ✓ [{result.get('response_code','?')}] "
                                  f"Fit:{result.get('fit_score',0)}/5  Risk:{result.get('risk_score',0)}/5",
                                  {"current": pct, "total": total})

            # ── 5. Scoring Agent ──────────────────────────────────────────────
            self.emit("agent_start", "Scoring Agent", "Computing overall fit and risk scores...")
            avg_fit = round(sum(fit_scores) / len(fit_scores), 2) if fit_scores else 0.0
            avg_risk = round(sum(risk_scores) / len(risk_scores), 2) if risk_scores else 0.0
            self.emit("agent_complete", "Scoring Agent",
                      f"Overall Fit: {avg_fit}/5  |  Overall Risk: {avg_risk}/5",
                      {"fit": avg_fit, "risk": avg_risk})

            # ── 6. Review Agent ───────────────────────────────────────────────
            self.emit("agent_start", "Review Agent", "Quality-checking responses and flagging high-risk items...")
            self._review_answers(rfp_id)
            self.emit("agent_complete", "Review Agent",
                      f"Review complete — {answered} answered, {flagged} flagged for human review")

            self.db.update_rfp(rfp_id,
                status="complete",
                fit_score=avg_fit,
                risk_score=avg_risk,
                question_count=len(questions),
                answered_count=answered,
                flagged_count=flagged,
                processed_at="datetime('now')",
            )

            self.emit("processing_complete", data={
                "rfp_id": rfp_id,
                "fit_score": avg_fit,
                "risk_score": avg_risk,
                "answered": answered,
                "flagged": flagged,
                "total": len(questions),
            })

        except Exception as e:
            msg = str(e)
            self.db.update_rfp(rfp_id, status="error", last_error=msg)
            self.emit("error", message=msg)
            raise

    def process_document(self, rfp_id: int, doc_id: int, filepath: str) -> None:
        """Process a single document within a multi-document RFP project."""
        self.db.update_document(doc_id, status="processing")
        self.emit("agent_start", "Parser Agent",
                  f"Reading document {doc_id}: {os.path.basename(filepath)}")
        try:
            questions = self._parse_rfp(rfp_id, filepath)
            if not questions:
                self.db.update_document(doc_id, status="error")
                fname = os.path.basename(filepath)
                ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                if ext == "docx":
                    self.emit("error", message=(
                        f"No requirements could be extracted from {fname}. "
                        "For Word documents, NaughtRFP works best with documents that contain "
                        "tables or numbered requirement lists. "
                        "Try converting to CSV or XLSX for best results."
                    ))
                else:
                    self.emit("error", message=f"No requirements found in {fname}.")
                return

            # Tag all questions with this document_id
            for q in questions:
                self.db.update_question(q["id"], document_id=doc_id)

            self.emit("agent_complete", "Parser Agent",
                      f"Extracted {len(questions)} requirements from {os.path.basename(filepath)}")

            questions = self._analyze_questions(rfp_id, questions)

            self.emit("agent_start", "Research Agent", "Pre-fetching KB context...")
            kb_cache: dict[int, list] = {}
            for q in questions:
                hits = self.db.search_knowledge_base(q["question_text"][:150], limit=3)
                if hits:
                    kb_cache[q["id"]] = hits

            self.emit("agent_start", "Answer Agent",
                      f"Processing {len(questions)} requirements ({_ANSWER_WORKERS} workers)...")

            answered = flagged = 0
            fit_scores: list[float] = []
            risk_scores: list[float] = []
            completed = [0]

            def _process_one(args):
                idx, q = args
                return idx, q, self._research_and_answer(
                    rfp_id, q, pre_context=kb_cache.get(q["id"], []))

            with ThreadPoolExecutor(max_workers=_ANSWER_WORKERS) as pool:
                futures = {pool.submit(_process_one, (i, q)): i
                           for i, q in enumerate(questions)}
                for future in as_completed(futures):
                    try:
                        _, q, result = future.result()
                    except Exception as e:
                        flagged += 1
                        self.emit("agent_progress", "Answer Agent", f"  ✗ Error: {e}")
                        continue
                    completed[0] += 1
                    if result.get("flagged"):
                        flagged += 1
                    else:
                        answered += 1
                        fit_scores.append(result.get("fit_score", 3))
                        risk_scores.append(result.get("risk_score", 3))
                    self.emit("agent_progress", "Answer Agent",
                              f"  [{completed[0]}/{len(questions)}] "
                              f"{'⚑' if result.get('flagged') else '✓'} "
                              f"{q['category'][:40]}",
                              {"current": completed[0], "total": len(questions)})

            self._review_answers(rfp_id)

            self.db.update_document(doc_id,
                status="complete",
                question_count=len(questions),
                answered_count=answered,
                flagged_count=flagged,
                processed_at="now",
            )
            self.db.sync_rfp_counts(rfp_id)

            self.emit("processing_complete", data={
                "doc_id": doc_id,
                "rfp_id": rfp_id,
                "answered": answered,
                "flagged": flagged,
                "total": len(questions),
            })

        except Exception as e:
            self.db.update_document(doc_id, status="error")
            self.emit("error", message=str(e))
            raise

    # ── Parser Agent ──────────────────────────────────────────────────────────

    # Sheet names that are almost certainly navigation/cover sheets — skip them
    _NAV_SHEET_PATTERNS = (
        "cover", "instruction", "readme", "overview", "index", "nav",
        "help", "legend", "key", "copyright", "summary", "toc", "table of",
        "how to", "guide", "notes", "ref", "drop", "lookup", "list",
    )

    def _parse_rfp(self, rfp_id, filepath):
        ext = filepath.rsplit(".", 1)[-1].lower()
        raw_rows: list[dict] = []

        if ext == "csv":
            with open(filepath, newline="", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    raw_rows.append(dict(row))

        elif ext in ("xlsx", "xls", "xlsm"):
            # keep_vba=True preserves macros so they survive round-trip export
            wb = openpyxl.load_workbook(filepath, data_only=True, keep_vba=(ext == "xlsm"))

            has_macros = (ext == "xlsm") or bool(getattr(wb, "vba_archive", None))
            sheets_found = wb.sheetnames

            if has_macros:
                self.emit("agent_progress", "Parser Agent",
                          f"📎 Macro-enabled workbook (.xlsm) — reading computed cell values across {len(sheets_found)} sheet(s)")

            # Decide which sheets to parse
            def _is_nav_sheet(name: str) -> bool:
                n = name.lower().strip()
                return any(p in n for p in self._NAV_SHEET_PATTERNS)

            def _sheet_has_content(ws) -> bool:
                """Return True if the sheet looks like it has requirement-style data."""
                if ws.max_row < 3 or ws.max_column < 2:
                    return False
                # Count non-empty cells in first 10 rows
                cells_with_text = sum(
                    1 for row in ws.iter_rows(min_row=1, max_row=10, values_only=True)
                    for c in row
                    if c and len(str(c).strip()) > 5
                )
                return cells_with_text >= 5

            candidate_sheets = [
                s for s in sheets_found
                if not _is_nav_sheet(s) and _sheet_has_content(wb[s])
            ]

            # If everything was skipped, fall back to all non-empty sheets
            if not candidate_sheets:
                candidate_sheets = [s for s in sheets_found if _sheet_has_content(wb[s])]
            # Final fallback: active sheet
            if not candidate_sheets:
                candidate_sheets = [wb.active.title]

            self.emit("agent_progress", "Parser Agent",
                      f"Scanning {len(candidate_sheets)}/{len(sheets_found)} sheet(s): {', '.join(candidate_sheets[:5])}")

            for sheet_name in candidate_sheets:
                ws = wb[sheet_name]
                headers = None
                for row in ws.iter_rows(values_only=True):
                    if not any(c for c in row if c is not None):
                        continue
                    if headers is None:
                        headers = [str(c).strip() if c is not None else f"col_{i}"
                                   for i, c in enumerate(row)]
                    else:
                        row_dict = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
                        row_dict["_sheet"] = sheet_name  # track source tab
                        raw_rows.append(row_dict)

        elif ext == "docx":
            # Extract content from Word document using python-docx.
            # Tables are the primary source (most RFPs use requirement tables);
            # paragraphs and full-text are fallbacks for unstructured documents.
            raw_rows = _parse_docx_rows(filepath)

        if not raw_rows:
            return []

        # Determine minimum text length: DOCX paragraph rows need a higher bar
        # than structured CSV/XLSX cells to filter noise.
        is_docx = ext == "docx"
        _min_text_len = 20 if is_docx else 10

        headers_list = list(raw_rows[0].keys())
        sample = json.dumps(
            [{k: str(v)[:120] for k, v in r.items()} for r in raw_rows[:4]],
            default=str,
        )

        resp = self.client.messages.create(
            model=_MODEL_FAST,
            max_tokens=512,
            system="You analyze RFP spreadsheet structures. Respond with valid JSON only, no markdown.",
            messages=[{                "role": "user",
                "content": f"""Headers: {headers_list}
Sample rows: {sample}

Identify columns. Respond with JSON:
{{"requirement_column": "exact header name with main requirement/question text",
  "category_column": "exact header for category/section or null",
  "priority_column": "exact header for priority/criticality or null",
  "response_column": "exact header where vendor response code goes or null",
  "comments_column": "exact header for vendor comments or null"}}"""
            }],
        )

        _record(self.db, rfp_id, "Parser Agent", resp)
        mapping = _extract_json(resp.content[0].text) or {}
        req_col = mapping.get("requirement_column")
        cat_col = mapping.get("category_column")
        pri_col = mapping.get("priority_column")

        if not req_col or req_col not in headers_list:
            # Fallback: use longest text column
            req_col = max(headers_list, key=lambda h: sum(
                len(str(r.get(h, "") or "")) for r in raw_rows[:10]
            ))

        # For DOCX rows from _parse_docx_rows, "requirement" and "text" are the
        # natural requirement columns.  If the Parser Agent missed them (e.g.
        # because it saw "requirement" as a generic label), override here.
        if is_docx and req_col not in headers_list:
            for preferred in ("requirement", "text"):
                if preferred in headers_list:
                    req_col = preferred
                    break

        # When DOCX rows carry a "category" column, prefer it over the Parser
        # Agent's detection (which may have returned null for paragraph-based docs).
        if is_docx and not cat_col and "category" in headers_list:
            cat_col = "category"

        # Numbered-item pattern for DOCX rows: don't filter these out even if
        # they are short — "1. Provide proof of insurance" is 34 chars but valid.
        _NUMBERED_PREFIX = re.compile(r'^(\(?\d+\)?\.?\)?\s|\(?[a-zA-Z]\)?\.?\)?\s)')

        questions = []
        seen = set()
        for i, row in enumerate(raw_rows):
            text = str(row.get(req_col, "") or "").strip()
            # For DOCX rows, also check the "requirement" key directly if req_col differs
            if not text and is_docx and "requirement" in row:
                text = str(row.get("requirement", "") or "").strip()
            if not text and is_docx and "text" in row:
                text = str(row.get("text", "") or "").strip()

            # Length check — relax for numbered items in DOCX
            is_numbered_item = is_docx and bool(_NUMBERED_PREFIX.match(text))
            effective_min = 10 if is_numbered_item else _min_text_len

            if (not text or len(text) < effective_min or
                    text.lower() in ("requirement", "criteria", "question", req_col.lower()) or
                    text in seen):
                continue
            seen.add(text)

            category = str(row.get(cat_col, "") or "").strip() if cat_col else "General"
            priority = str(row.get(pri_col, "") or "").strip() if pri_col else ""
            sheet    = str(row.get("_sheet", "") or "").strip()

            # Prefix category with sheet name when multiple sheets are present
            # so the SE knows which tab each requirement came from
            if sheet and sheet not in (category, ""):
                display_category = f"{sheet} › {category}" if category and category != "General" else sheet
            else:
                display_category = category or "General"

            q_id = self.db.create_question(rfp_id, i, display_category, text)
            questions.append({
                "id": q_id,
                "row_index": i,
                "category": display_category,
                "question_text": text,
                "priority": priority,
                "column_mapping": mapping,
                "source_sheet": sheet,
            })

        return questions

    # ── Analysis Agent ────────────────────────────────────────────────────────

    def _analyze_questions(self, rfp_id, questions):
        questions_text = "\n".join(
            f"{i+1}. [{q['category']}] {q['question_text'][:180]}"
            for i, q in enumerate(questions)
        )

        resp = self.client.messages.create(
            model=_MODEL,
            max_tokens=4096,
            system="You are an expert Okta SE analyzing RFP requirements. Respond with a JSON array only, no markdown.",
            messages=[{
                "role": "user",
                "content": f"""Analyze these RFP requirements and return a JSON array:

{questions_text}

Return: [{{"index":1,"okta_products":["OIG","LCM"],"refined_category":"Identity Lifecycle","risk_score":2}}]

risk_score: 1=low, 5=high (5=legal/SLA/pricing commitments)
okta_products options: OIG, LCM, Workflows, SSO, MFA, Universal Directory, PAM, AI, Access Gateway, OIN"""
            }],
        )

        _record(self.db, rfp_id, "Analysis Agent", resp)
        analysis = _extract_json(resp.content[0].text, kind="array") or []
        analysis_map = {a["index"]: a for a in analysis if isinstance(a, dict)}

        for i, q in enumerate(questions):
            a = analysis_map.get(i + 1, {})
            q["okta_products"] = a.get("okta_products", [])
            q["refined_category"] = a.get("refined_category", q["category"])
            q["risk_score"] = a.get("risk_score", 3)
            self.db.update_question(
                q["id"],
                category=q["refined_category"],
                okta_products=json.dumps(q["okta_products"]),
                risk_score=q["risk_score"],
            )

        return questions

    # ── Research + Answer Agent ───────────────────────────────────────────────

    def _research_and_answer(self, rfp_id: int, q: dict, pre_context: list | None = None) -> dict:
        # Build pre-loaded KB context string to inject directly (eliminates 1 tool round-trip)
        kb_hint = ""
        if pre_context:
            entries = [{"q": r["question"][:120], "a": r["answer"][:250], "src": r.get("source_rfp_name", "")}
                       for r in pre_context]
            kb_hint = f"\n\nPre-loaded KB matches (use these if relevant before calling tools):\n{json.dumps(entries, separators=(',', ':'))}"

        system = (
            "You are a senior Okta Solutions Engineer responding to vendor RFP requirements.\n"
            f"{OKTA_KNOWLEDGE}\n"
            "RULES:\n"
            "- Represent Okta accurately. Response codes: F=Full, P=Partial, C=Custom, NE=Planned, N=Not Available.\n"
            "- Use pre-loaded KB matches if they address the requirement; otherwise call search_knowledge_base.\n"
            "- If KB is insufficient, call search_web for current Okta data.\n"
            "- Flag for review if confidence<0.60 or requirement involves pricing/SLA/legal commitments.\n"
            "- Answer in 2-4 sentences. Cite the Okta product(s) used."
        )

        user_content = (
            f"Category: {q['refined_category']} | Priority: {q.get('priority','?')} | "
            f"Products: {', '.join(q.get('okta_products', []))}\n\n"
            f"REQUIREMENT:\n{q['question_text']}"
            f"{kb_hint}\n\n"
            "Use KB context above if relevant, otherwise call search_knowledge_base.\n"
            "Either call flag_for_review OR return JSON:\n"
            '{"response_code":"F|P|C|NE|N","answer":"2-4 sentences","confidence":0.0-1.0,'
            '"fit_score":1-5,"risk_score":1-5,"sources":["source"],"okta_products":["product"]}'
        )

        messages = [{"role": "user", "content": user_content}]

        flag_reason = None
        final_text = None

        for _ in range(3):  # max 3 iterations: pre-context already loaded, fewer round trips needed
            resp = self.client.messages.create(
                model=_MODEL,
                max_tokens=768,  # answers are short; 768 is sufficient
                system=system,
                messages=messages,
                tools=([SEARCH_WEB_TOOL] if _WEB_SEARCH_ENABLED else []) + [SEARCH_KB_TOOL, FLAG_REVIEW_TOOL],
            )

            if resp.stop_reason == "tool_use":
                messages.append({"role": "assistant", "content": resp.content})
                tool_results = []

                for block in resp.content:
                    if block.type != "tool_use":
                        continue
                    if block.name == "search_web":
                        query = block.input.get("query", "")
                        self.emit("agent_progress", "Research Agent",
                                  f"🌐 Web search: \"{query[:70]}\"")
                        web_result = _do_web_search(
                            query,
                            self.api_key,
                            _LITELLM_BASE_URL,
                        )
                        tool_results.append({
                            "type": "tool_result",
                            "tool_use_id": block.id,
                            "content": web_result,
                        })
                    elif block.name == "search_knowledge_base":
                        kb = self.db.search_knowledge_base(block.input.get("query", ""), limit=3)
                        result_text = json.dumps([
                            {"question": r["question"][:150], "answer": r["answer"][:300],
                             "category": r["category"]}
                            for r in kb
                        ])
                        tool_results.append({
                            "type": "tool_result",
                            "tool_use_id": block.id,
                            "content": result_text if result_text != "[]" else "No matching entries found.",
                        })
                    elif block.name == "flag_for_review":
                        flag_reason = block.input.get("reason", "Needs human review")
                        tool_results.append({
                            "type": "tool_result",
                            "tool_use_id": block.id,
                            "content": "Question has been flagged for human review.",
                        })

                messages.append({"role": "user", "content": tool_results})

                if flag_reason:
                    break

            else:
                _record(self.db, rfp_id, "Answer Agent", resp)
                for block in resp.content:
                    if hasattr(block, "text"):
                        final_text = block.text
                break

        if flag_reason:
            self.db.update_question(q["id"],
                status="flagged", needs_review=1, review_reason=flag_reason)
            return {"flagged": True, "review_reason": flag_reason}

        data = _extract_json(final_text or "") or {}

        if not data.get("answer"):
            self.db.update_question(q["id"],
                status="flagged", needs_review=1,
                review_reason="Could not generate a structured response")
            return {"flagged": True, "review_reason": "Parse error"}

        self.db.update_question(q["id"],
            answer=data["answer"],
            response_code=data.get("response_code", "P"),
            confidence=data.get("confidence", 0.7),
            fit_score=data.get("fit_score", 3),
            risk_score=data.get("risk_score", 3),
            sources=json.dumps(data.get("sources", [])),
            okta_products=json.dumps(data.get("okta_products", q.get("okta_products", []))),
            status="answered",
        )

        return data

    # ── Review Agent ──────────────────────────────────────────────────────────

    def _review_answers(self, rfp_id):
        questions = self.db.get_questions(rfp_id)
        for q in questions:
            if q["status"] == "answered" and q.get("risk_score", 0) >= 4:
                answer = q.get("answer", "")
                if answer and "legal" not in answer.lower() and "commercial" not in answer.lower():
                    note = ("\n\n⚠ This requirement has been scored as high-risk. "
                            "Review with your legal/commercial team before submission.")
                    self.db.update_question(q["id"], answer=answer + note)


# ── Knowledge Base Ingestion Agent ────────────────────────────────────────────

class KnowledgeBaseAgent:
    def __init__(self, api_key, db, event_queue):
        self.api_key = api_key
        self.db = db
        self.q = event_queue

    def emit(self, event_type, message="", data=None):
        self.q.put({
            "type": event_type,
            "agent": "KB Ingestion Agent",
            "message": message,
            **({"data": data} if data else {}),
            "timestamp": time.time(),
        })

    def ingest_rfp(self, rfp_id):
        rfp = self.db.get_rfp(rfp_id)
        questions = self.db.get_questions(rfp_id)
        answered = [q for q in questions if q["status"] == "answered" and q.get("answer")]

        self.emit("agent_start",
                  f'Ingesting {len(answered)} answered requirements from "{rfp["name"]}"...')

        ingested = 0
        for q in answered:
            existing = self.db.search_knowledge_base(q["question_text"][:80], limit=1)
            if existing:
                continue
            self.db.add_to_knowledge_base(
                source_rfp_name=rfp["name"],
                category=q["category"],
                question=q["question_text"],
                answer=q["answer"],
                response_code=q.get("response_code"),
                okta_products=json.loads(q.get("okta_products") or "[]"),
            )
            ingested += 1

        self.emit("agent_complete",
                  f"Added {ingested} new entries to the knowledge base",
                  {"ingested": ingested})
        return ingested


# ── Demo Prep Agent ───────────────────────────────────────────────────────────

class DemoPrepAgent:
    def __init__(self, api_key, db, event_queue):
        self.api_key = api_key
        self.db = db
        self.q = event_queue
        self.client = _make_client(api_key)

    def emit(self, event_type, message="", data=None):
        self.q.put({
            "type": event_type,
            "agent": "Demo Prep Agent",
            "message": message,
            **({"data": data} if data else {}),
            "timestamp": time.time(),
        })

    def generate(self, rfp_id, customer_format=None):
        rfp = self.db.get_rfp(rfp_id)
        questions = self.db.get_questions(rfp_id)
        answered = [q for q in questions if q["status"] == "answered" and q.get("answer")]
        flagged  = [q for q in questions if q["status"] == "flagged"]

        if not answered:
            self.emit("error", "No answered questions — process the RFP first.")
            return None

        ci = json.loads(rfp.get("customer_info") or "{}") if rfp.get("customer_info") else {}
        customer = ci.get("customer_name", rfp["name"])
        industry = ci.get("industry", "")
        scope    = ci.get("scope_summary", "")

        self.emit("agent_start",
                  f"Analysing {len(answered)} answered requirements for {customer}…")

        # Build compact question digest for the prompt
        q_digest = []
        for q in answered:
            products = json.loads(q.get("okta_products") or "[]")
            q_digest.append({
                "id": q["id"],
                "category": q["category"],
                "requirement": q["question_text"][:150],
                "response_code": q.get("response_code", "F"),
                "fit": q.get("fit_score", 3),
                "risk": q.get("risk_score", 3),
                "okta_products": products,
            })

        format_instruction = ""
        if customer_format:
            format_instruction = (
                f"\n\nThe customer has provided this required demo format/agenda:\n"
                f"{customer_format[:1500]}\n"
                "Map Okta capabilities to each section they require."
            )

        flagged_note = ""
        if flagged:
            flagged_note = (
                f"\n\nNOTE: {len(flagged)} requirements were flagged for human review. "
                "Include a 'Questions to Address' section at the end covering these gaps."
            )

        self.emit("agent_progress", "Demo Prep Agent",
                  "Building APEX brief (Before/After/PBOs/Required Capabilities)…")

        resp = self.client.messages.create(
            model=_MODEL,
            max_tokens=8000,
            system=(
                "You are a senior Okta Solutions Engineer expert in the APEX sales framework "
                "(Okta's implementation of Command of the Message + MEDDPICCC). "
                "Respond with valid compact JSON only — no markdown fences, no prose."
            ),
            messages=[{
                "role": "user",
                "content": f"""Create an APEX-aligned demo preparation plan for this RFP. Return ONLY compact JSON.

Customer: {customer} | Industry: {industry}
Scope: {scope}

Requirements ({len(answered)} answered, {len(flagged)} flagged):
{json.dumps(q_digest[:30], separators=(',', ':'))}
{format_instruction}
{flagged_note}

Return this exact JSON structure (all strings max 30 words):
{{
  "apex_brief": {{
    "mantra": "1-2 sentence power statement: customer name + their After Scenario + key PBO in customer language",
    "before_scenario": "Current state pain — what is broken/manual/risky today for this customer",
    "negative_consequences": ["Quantified cost/risk of staying in the Before Scenario (2-3 items)"],
    "after_scenario": "Desired future state — what success looks like with Okta",
    "positive_business_outcomes": ["Measurable outcome 1 (with metric)", "Measurable outcome 2", "Measurable outcome 3"],
    "required_capabilities": ["Capability framed as objectively necessary (maps to Okta strength)"],
    "unique_differentiators": ["Why Okta specifically — not just any IGA vendor"]
  }},
  "executive_summary": "2 sentence SE-perspective summary of what to demo and why it wins",
  "total_minutes": 60,
  "sections": [
    {{
      "order": 1,
      "title": "Section title",
      "okta_products": ["OIG"],
      "priority": "critical|high|medium",
      "estimated_minutes": 15,
      "requirement_ids": [1, 2],
      "pbo_addressed": "Which PBO this section proves",
      "required_capability": "Which Required Capability this demonstrates",
      "demo_scenario": "1 sentence customer-language framing for this section",
      "demo_steps": ["Step 1: Show X", "Step 2: Demonstrate Y", "Step 3: Prove Z"],
      "com_talking_points": ["CoM-framed point using Before/After language", "Differentiator point"],
      "differentiators": ["Specific Okta advantage vs competition here"]
    }}
  ],
  "questions_to_address": ["Flagged/uncertain area to prepare for"],
  "recommended_demo_env": "What to configure in the demo org before the call",
  "discovery_questions": ["Question to ask in pre-demo discovery to sharpen the After Scenario"]
}}

Rules:
- APEX brief must come from the RFP content — infer Before/After from what the customer is asking to fix
- PBOs must be measurable (include numbers/percentages where possible)
- Required Capabilities should favor Okta differentiators as objectively necessary
- 4-6 demo sections, ordered by customer priority (Critical first)
- Each section must map to at least one PBO
- CoM talking points use Before → After language, not feature lists
- Discovery questions help the SE validate/sharpen the APEX brief before the demo"""
            }],
        )
        _record(self.db, rfp_id, "Demo Prep Agent", resp)

        raw = resp.content[0].text
        if resp.stop_reason == "max_tokens":
            self.emit("error", "Response truncated — too many requirements. Try a smaller RFP.")
            return None

        data = _extract_json(raw) or {}
        if not data.get("sections"):
            preview = raw[:200].replace('\n', ' ')
            self.emit("error", f"Could not parse demo plan. Model returned: {preview}")
            return None

        sections  = data["sections"]
        total_min = data.get("total_minutes", sum(s.get("estimated_minutes", 10) for s in sections))
        summary   = data.get("executive_summary", "")
        apex      = data.get("apex_brief", {})

        mantra = apex.get("mantra", "")
        self.emit("agent_complete",
                  f"APEX brief + {len(sections)}-section demo plan ({total_min} min) — \"{mantra[:80]}\"",
                  {"sections": len(sections), "total_minutes": total_min})

        plan_id = self.db.create_demo_plan(rfp_id, sections, summary, total_min)
        self.db.update_demo_plan(plan_id,
            notes=json.dumps({
                "apex_brief":            apex,
                "questions_to_address":  data.get("questions_to_address", []),
                "recommended_demo_env":  data.get("recommended_demo_env", ""),
                "discovery_questions":   data.get("discovery_questions", []),
            })
        )
        return plan_id


# ── AI Search ─────────────────────────────────────────────────────────────────

def ai_search_knowledge_base(api_key: str, db, query: str) -> dict:
    """
    Search the KB with AI assistance.
    Returns: {"bluf": "...", "results": [...], "query": "..."}
    """
    client  = _make_client(api_key)
    entries = db.get_kb_entries(limit=100)

    # Always include FTS fallback results so we have something to BLUF
    fts_results = db.search_knowledge_base(query, limit=8)

    if not entries:
        return {"bluf": None, "results": fts_results, "query": query}

    entries_text = "\n".join(
        f"{i+1}. [{e['category']}] Q: {e['question'][:110]} | A: {e['answer'][:120]}"
        for i, e in enumerate(entries[:60])
    )

    resp = client.messages.create(
        model=_MODEL,
        max_tokens=900,
        system=(
            "You are an expert Okta Solutions Engineer with deep knowledge of identity, "
            "security, and compliance. You help SEs find relevant knowledge base entries.\n"
            "Respond with valid JSON only, no markdown fences."
        ),
        messages=[{
            "role": "user",
            "content": (
                f'Search query: "{query}"\n\n'
                "Rules for matching:\n"
                "- Be LIBERAL: match acronyms (DR=disaster recovery, BCP=business continuity, "
                "MFA=multi-factor auth, SoD=separation of duties, LCM=lifecycle management, "
                "PAM=privileged access, SSO=single sign-on, IGA=identity governance, "
                "NHI=non-human identity, OIG=Okta Identity Governance, PIPEDA=Canadian privacy law).\n"
                "- Match fragments and partial phrases — 'uptime' should match '99.99% SLA'.\n"
                "- Match by CONTEXT and INTENT, not just exact words.\n"
                "- A single word query like 'encryption' should match entries about AES, TLS, key management, etc.\n"
                "- A sentence query should find entries that address ANY part of it.\n\n"
                f"Knowledge base entries:\n{entries_text}\n\n"
                "Tasks:\n"
                "1. Return indices of the 1-10 most relevant entries (liberal matching).\n"
                "2. Write a BLUF (Bottom Line Up Front): 2-4 sentences stating Okta's position "
                "on this topic. Military style — most important fact first. Include key numbers, "
                "certifications, or caveats. If entries show Okta lacks something, say so clearly.\n\n"
                'JSON: {"indices":[1,3,7],"explanations":["brief why each is relevant"],'
                '"bluf":"Okta... concise synthesis."}'
            )
        }],
    )

    data         = _extract_json(resp.content[0].text) or {}
    indices      = data.get("indices", [])
    explanations = data.get("explanations", [])
    bluf         = data.get("bluf", "")

    results = []
    for i, idx in enumerate(indices):
        if 1 <= idx <= len(entries):
            entry = dict(entries[idx - 1])
            entry["relevance"] = explanations[i] if i < len(explanations) else ""
            entry["ai_match"]  = True
            results.append(entry)

    if not results:
        results = fts_results

    return {"bluf": bluf or None, "results": results, "query": query}


# ── Customer Detection Agent ──────────────────────────────────────────────────

def _sample_file_text(filepath, max_chars=3000):
    """Extract a text sample from a CSV, XLSX, or DOCX for customer detection."""
    ext = filepath.rsplit(".", 1)[-1].lower()
    chunks = []

    try:
        if ext == "csv":
            with open(filepath, encoding="utf-8-sig", errors="ignore") as f:
                chunks.append(f.read(max_chars))
        elif ext in ("xlsx", "xls"):
            wb = openpyxl.load_workbook(filepath, data_only=True)
            for sheet in wb.worksheets[:2]:
                for row in sheet.iter_rows(values_only=True):
                    line = " | ".join(str(c) for c in row if c is not None)
                    if line.strip():
                        chunks.append(line)
                    if sum(len(c) for c in chunks) > max_chars:
                        break
                if sum(len(c) for c in chunks) > max_chars:
                    break
        elif ext == "docx":
            try:
                from docx import Document
                doc = Document(filepath)
                # Sample first two tables + first 20 paragraphs
                for table in doc.tables[:2]:
                    for row in table.rows:
                        line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                        if line:
                            chunks.append(line)
                        if sum(len(c) for c in chunks) > max_chars:
                            break
                for para in doc.paragraphs[:20]:
                    text = para.text.strip()
                    if text:
                        chunks.append(text)
                    if sum(len(c) for c in chunks) > max_chars:
                        break
            except ImportError:
                pass
    except Exception:
        pass

    return "\n".join(chunks)[:max_chars]


def detect_customer(api_key, filepath, rfp_filename):
    """Use Claude to identify who the RFP is for and what it covers."""
    client = _make_client(api_key)
    sample = _sample_file_text(filepath)

    resp = client.messages.create(
        model=_MODEL_FAST,
        max_tokens=512,
        system="You are an expert at reading RFP documents. Respond with valid JSON only, no markdown.",
        messages=[{
            "role": "user",
            "content": f"""Analyze this RFP file and extract key identifying information.

Filename: {rfp_filename}
Content sample:
{sample}

Respond with JSON:
{{
  "customer_name": "Full organization name issuing the RFP",
  "customer_short": "Short name or abbreviation (e.g. 'AHS', 'SaskPower')",
  "rfp_number": "Reference/procurement number if present, else null",
  "project_name": "What this RFP is for (2-6 words)",
  "industry": "Sector (e.g. Healthcare, Energy, Government, Financial Services, Food & Beverage)",
  "scope_summary": "1-2 sentence plain-English summary of what they are buying",
  "issuing_department": "Department or team name if mentioned, else null",
  "estimated_scale": "Scale of deployment if mentioned (e.g. '150,000 identities'), else null",
  "confidence": 0.0-1.0
}}"""
        }],
    )

    data = _extract_json(resp.content[0].text) or {}
    if not data.get("customer_name"):
        # Fallback: derive from filename
        name = rfp_filename.replace("_", " ").replace("-", " ").split(".")[0]
        data = {
            "customer_name": name,
            "customer_short": name[:20],
            "rfp_number": None,
            "project_name": "RFP",
            "industry": "Unknown",
            "scope_summary": "No additional context available.",
            "issuing_department": None,
            "estimated_scale": None,
            "confidence": 0.2,
        }
    return data
